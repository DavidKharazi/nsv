
import re
from datetime import datetime
from typing import Optional, Dict, Any, List, Tuple
from langchain_core.documents import Document
from langchain_community.vectorstores import Chroma
from langchain_openai import ChatOpenAI, OpenAIEmbeddings, OpenAI
import sqlite3
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.prompts.chat import ChatPromptTemplate, MessagesPlaceholder
from langchain.schema import AIMessage, HumanMessage
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException, Form, Request, status, Depends
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse, JSONResponse
import uvicorn
import boto3
from fnmatch import fnmatchcase
import json
import os
from docx.oxml.ns import qn
from docx import Document as DocxDocument
from io import BytesIO
from authlib.integrations.starlette_client import OAuth
from dotenv import load_dotenv
import uuid
from sqlalchemy.orm import sessionmaker
from sqlalchemy import create_engine, text
import requests
from requests import Session



load_dotenv()


os.environ['OPENAI_API_KEY'] = 'my_api'


model_name = "gpt-4o"
temperature = 0
llm = ChatOpenAI(model=model_name, temperature=temperature)
embeddings = OpenAIEmbeddings()

current_user = 'nsv'

# Настройка клиента для Yandex S3
session = boto3.session.Session()
s3_client = session.client(
    service_name='s3',
    endpoint_url='https://storage.yandexcloud.net',
    aws_access_key_id='my_aws',
    aws_secret_access_key='my_secret',
)

CHROMA_PATH = f'./chroma/{current_user}/'

oauth = OAuth()


def init_metadata_db():
    with sqlite3.connect('metadata.db') as conn:
        conn.execute('''
            CREATE TABLE IF NOT EXISTS Admin (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                email TEXT UNIQUE NOT NULL,
                password TEXT NOT NULL
            )
            ''')
        conn.execute('''
                CREATE TABLE IF NOT EXISTS uploaded_docs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                global_source TEXT,
                filename TEXT
                );
                ''')
        conn.execute('''
        CREATE TABLE IF NOT EXISTS Users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email VARCHAR(255),
            name VARCHAR(255),
            cyberman_id INTEGER,
            chat_id INTEGER,
            is_active BOOLEAN DEFAULT FALSE,
            confirmation_token TEXT,
            reset_token TEXT,
            new_password VARCHAR(255),
            FOREIGN KEY (cyberman_id) REFERENCES Cyberman(id),
            FOREIGN KEY (chat_id) REFERENCES Chat(id)
        );
        ''')
        conn.execute('''
        CREATE TABLE IF NOT EXISTS Cyberman (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name VARCHAR(255),
            creativity DOUBLE,
            prompt VARCHAR(255)
        );
        ''')
        conn.execute('''
        CREATE TABLE IF NOT EXISTS Session (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            started_at TIMESTAMP,
            ended_at TIMESTAMP,
            user_id INTEGER,
            cyberman_id INTEGER,
            topic TEXT,
            FOREIGN KEY (user_id) REFERENCES Users(id),
            FOREIGN KEY (cyberman_id) REFERENCES Cyberman(id)
        );
        ''')
        conn.execute('''
        CREATE TABLE IF NOT EXISTS Chat (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            messages VARCHAR(255),
            sender VARCHAR(255),
            sent_at TIMESTAMP DEFAULT (datetime('now', 'localtime')),
            session_id INTEGER,
            FOREIGN KEY (session_id) REFERENCES Session(id)
        );
        ''')


init_metadata_db()


class DatabaseManager:
    def __init__(self, db_path="metadata.db"):
        self.connection = sqlite3.connect(db_path)
        self.connection.row_factory = sqlite3.Row  # Позволяет обращаться к колонкам по именам
        self.cursor = self.connection.cursor()
        self.db_path = db_path
        self.connection.row_factory = sqlite3.Row  # Позволяет работать с результатами в виде объектов Row

    def get_chats_by_user_id(self, user_id):
        query = """
        SELECT id, cyberman_id, started_at FROM Session WHERE user_id = ?
        """
        cursor = self.connection.cursor()
        cursor.execute(query, (user_id,))
        rows = cursor.fetchall()
        return [dict(row) for row in rows]

    def create_new_chat_session(self, user_id, cyberman_id):
        query = """
            INSERT INTO Session (user_id, cyberman_id) VALUES (?, ?)
            """
        cursor = self.connection.cursor()
        cursor.execute(query, (user_id, cyberman_id))
        self.connection.commit()
        return cursor.lastrowid

    def get_chat_messages_by_session_id(self, session_id):
        query = """
            SELECT sender, messages, sent_at FROM Chat WHERE session_id = ?
            """
        cursor = self.connection.cursor()
        cursor.execute(query, (session_id,))
        rows = cursor.fetchall()
        return [dict(row) for row in rows]

    def get_or_create_user(self, email, name, cyberman_id):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id FROM Users WHERE email = ?", (email,))
            user = cursor.fetchone()
            if user:
                print(f"Existing user found: {user[0]}")
                cursor.execute("UPDATE Users SET cyberman_id = ? WHERE email = ?", (cyberman_id, email))
                conn.commit()
                return user[0]
            else:
                confirmation_token = str(uuid.uuid4())
                cursor.execute(
                    "INSERT INTO Users (email, name, cyberman_id, is_active, confirmation_token) VALUES (?, ?, ?, ?, ?)",
                    (email, name, cyberman_id, False, confirmation_token)
                )
                user_id = cursor.lastrowid
                print(f"New user created: {user_id}")
                return user_id, confirmation_token

    def get_or_create_cyberman(self, name, creativity, prompt):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id FROM Cyberman WHERE name = ?", (name,))
            cyberman = cursor.fetchone()
            if cyberman:
                print(f"Existing cyberman found: {cyberman[0]}")
                return cyberman[0]
            else:
                cursor.execute("INSERT INTO Cyberman (name, creativity, prompt) VALUES (?, ?, ?)",
                               (current_user, creativity, prompt))
                cyberman_id = cursor.lastrowid
                print(f"New cyberman created: {cyberman_id}")
                return cyberman_id

    def create_session(self, user_id, cyberman_id):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            started_at = datetime.now()
            cursor.execute("INSERT INTO Session (started_at, user_id, cyberman_id) VALUES (?, ?, ?)",
                           (started_at, user_id, cyberman_id))
            session_id = cursor.lastrowid
            print(f"New session created: {session_id}")
            return session_id

    def end_session(self, session_id):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            ended_at = datetime.now()
            cursor.execute("UPDATE Session SET ended_at = ? WHERE id = ?", (ended_at, session_id))

    def add_chat_message(self, session_id, message, sender):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO Chat (messages, sender, session_id) VALUES (?, ?, ?)",
                           (message, sender, session_id))

    def get_session_by_user_and_cyberman(self, user_id, cyberman_id):
        query = """
        SELECT id FROM Session WHERE user_id = ? AND cyberman_id = ?
        """
        self.cursor.execute(query, (user_id, cyberman_id))
        return self.cursor.fetchone()

    def delete_chat(self, session_id: int):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            # Удаляем сообщения чата
            cursor.execute("DELETE FROM Chat WHERE session_id = ?", (session_id,))
            # Удаляем саму сессию
            cursor.execute("DELETE FROM Session WHERE id = ?", (session_id,))
            conn.commit()


db_manager = DatabaseManager()


class SQLiteChatHistory():
    def __init__(self, db_path="metadata.db", user_email=None, user_password=None, cyberman_name=None):
        self.db_path = db_path
        self.db_manager = DatabaseManager(db_path)
        self.current_session_id = None
        self.user_email = user_email
        self.user_password = user_password
        self.cyberman_name = cyberman_name

    def start_new_session(self, user_email=None, user_password=None, cyberman_name=current_user):
        user_email = user_email or self.user_email
        user_password = user_password or self.user_password
        cyberman_name = cyberman_name or self.cyberman_name

        print(f"Starting new session with: email={user_email}, cyberman={cyberman_name}")


        # Получаем или создаем Cyberman и получаем его ID
        cyberman_id = self.db_manager.get_or_create_cyberman(cyberman_name, temperature, prompt_sys)
        print(f"Cyberman ID: {cyberman_id}")

        # Передаем cyberman_id в метод get_or_create_user
        user_id = self.db_manager.get_or_create_user(user_email, user_password, cyberman_id)
        print(f"User ID: {user_id}")

        # Проверяем, существует ли уже сессия для данного пользователя и Cyberman
        session = self.db_manager.get_session_by_user_and_cyberman(user_id, cyberman_id)
        if session:
            self.current_session_id = session[0]
        else:
            self.current_session_id = self.db_manager.create_session(user_id, cyberman_id)
        print(f"Session ID: {self.current_session_id}")

        return self.current_session_id

    def add_message(self, message):
        if not self.current_session_id:
            print("No active session. Starting a new one.")
            self.start_new_session()

        print(f"Adding message to session {self.current_session_id}")

    def messages(self, limit=15):
        if not self.current_session_id:
            return ChatMessageHistory()

        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        c.execute(f"SELECT * FROM Chat WHERE session_id = ? ORDER BY id DESC LIMIT ?",
                  (self.current_session_id, limit))
        resp = c.fetchall()[::-1]
        chat_history = []
        for row in resp:
            id, message, sender, sent_at, session_id = row
            if sender == "human":
                chat_history.append(HumanMessage(content=message))
            elif sender == "ai":
                chat_history.append(AIMessage(content=message))
        conn.close()
        return ChatMessageHistory(messages=chat_history)

    def end_session(self):
        if self.current_session_id:
            self.db_manager.end_session(self.current_session_id)
            self.current_session_id = None


chat_history = SQLiteChatHistory()


def add_user_to_db(email: str, name: str, cyberman_id: int = None, chat_id: int = None):

    try:
        with sqlite3.connect('metadata.db') as conn:
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO Users (email, name, cyberman_id, chat_id, is_active) VALUES (?, ?, ?, ?, ?)",
                (email, name, cyberman_id, chat_id, False)
            )
            conn.commit()
            user_id = cursor.lastrowid  # Получаем ID добавленного пользователя
            return user_id

    except sqlite3.IntegrityError:
        raise HTTPException(status_code=400, detail="User already registered")



def delete_chat_history_last_n(self, n=10):
    conn = sqlite3.connect(self.db_path)
    c = conn.cursor()
    c.execute(f'''
    with max_id as (select max(id) as maxid from history_messages where user_id = '{current_user}')
    DELETE FROM history_messages
    WHERE id BETWEEN (select maxid from max_id) - {n} AND (select maxid from max_id)
    ''')
    conn.commit()
    conn.close()


def add_filename_to_metadata(source, filename):
    with sqlite3.connect('metadata.db') as conn:
        conn.execute(f'''INSERT INTO uploaded_docs (global_source, filename) values ('{source}', '{filename}') ; ''')


def delete_filename_from_metadata(source, filename):
    with sqlite3.connect('metadata.db') as conn:
        conn.execute(f'''DELETE from uploaded_docs where global_source = '{source}' and filename ='{filename}' ; ''')


class Document:
    def __init__(self, source: str, page_content: str, metadata: Optional[Dict[str, Any]] = None):
        self.source = source
        self.page_content = page_content
        self.metadata = metadata if metadata is not None else {'source': source}


def get_uploaded_filenames(source) -> List[str]:
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute(f"SELECT filename FROM uploaded_docs WHERE global_source = ?", (source,))
        rows = cursor.fetchall()
    filenames = [row[0] for row in rows]
    return filenames


def load_s3_files(bucket: str, prefix: str, suffix: str) -> List[str]:
    """List files in a given S3 bucket with a specified prefix and suffix."""
    try:
        response = s3_client.list_objects_v2(Bucket=bucket, Prefix=prefix)
        files = [content['Key'] for content in response.get('Contents', []) if content['Key'].endswith(suffix)]
        if not files:
            print(f"No files found in bucket {bucket} with prefix {prefix} and suffix {suffix}")
        else:
            print(f"Files found in bucket {bucket} with prefix {prefix} and suffix {suffix}: {files}")
        return files
    except Exception as e:
        print(f"Error listing files in bucket {bucket} with prefix {prefix} and suffix {suffix}: {e}")
        return []


def load_docx_new(source, bucket: str) -> List[Document]:
    prefix = 'nsv/docx/'
    suffix = '.docx'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    docs = []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = obj['Body'].read()

                # Используем BytesIO для чтения содержимого файла как бинарного потока
                doc_stream = BytesIO(content)
                doc = DocxDocument(doc_stream)

                # Извлекаем текст из документа docx
                full_text = []
                image_counter = 1

                # Получаем имя файла без расширения и создаем соответствующую папку
                filename_without_extension = os.path.splitext(os.path.basename(file))[0]
                image_folder = filename_without_extension  # Используем оригинальное имя файла для папки

                for para in doc.paragraphs:
                    # Обработка параграфов для создания ссылок на изображения
                    para_text = para.text
                    for run in para.runs:
                        for drawing in run.element.findall('.//a:blip', namespaces={
                            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                            image_rId = drawing.get(qn('r:embed'))
                            image_part = doc.part.related_parts[image_rId]
                            image_filename = f'image_{image_counter:02d}.{image_part.content_type.split("/")[-1]}'
                            image_counter += 1

                            # Загрузка изображения в бакет Яндекса
                            img_content = image_part.blob
                            s3_image_key = f"nsv/images/{image_folder}/{image_filename}"
                            s3_client.put_object(
                                Bucket=bucket,
                                Key=s3_image_key,
                                Body=img_content,
                                ContentDisposition='inline',
                                ContentType=image_part.content_type
                            )

                            # Генерация URL для изображения
                            s3_image_url = f"https://storage.yandexcloud.net/{bucket}/{s3_image_key}"
                            para_text += f'\n{s3_image_url}'
                    full_text.append(para_text)
                content = '\n'.join(full_text)

                docs.append(Document(source=file, page_content=content))
            except Exception as e:
                print(f"Error reading docx file {file}: {e}")

    return docs if docs else None


def load_txts(source, bucket: str) -> List[Document]:
    prefix = f'{current_user}/txt/'
    suffix = '.txt'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    docs = []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = obj['Body'].read().decode('utf-8')
                docs.append(Document(source=file, page_content=content))
            except Exception as e:
                print(f"Error reading txt file {file}: {e}")

    return docs if docs else None


def load_jsons(source, bucket: str) -> Tuple[List[Document], List[dict]]:
    prefix = f'{current_user}/json/'
    suffix = '.json'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    json_docs, json_metadata = [], []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = json.loads(obj['Body'].read().decode('utf-8'))
                json_docs.append(content)
                json_metadata.append({'source': file})
            except Exception as e:
                print(f"Error reading json file {file}: {e}")

    return (json_docs, json_metadata) if json_docs else (None, None)


def load_documents(global_source, bucket: str, file_types: List[str]) -> dict:
    """
    Загружаем документы в зависимости от типа документа из Yandex S3
    """
    all_docs = {'txt': None, 'json': None, 'json_metadata': None, 'docx': None}
    if 'txt' in file_types:
        txt_docs = load_txts(global_source, bucket)
        all_docs['txt'] = txt_docs
    if 'json' in file_types:
        json_docs, json_metadata = load_jsons(global_source, bucket)
        all_docs['json'] = json_docs
        all_docs['json_metadata'] = json_metadata
    if 'docx' in file_types:
        docx_docs = load_docx_new(global_source, bucket)
        all_docs['docx'] = docx_docs
    return all_docs


# Пример использования
DATA_BUCKET = 'utlik'
DOCS = load_documents('s3', DATA_BUCKET, ['txt', 'json', 'docx'])


def split_docs_to_chunks(documents: dict, file_types: List[str], keyword="—BLOCK—"):
    all_chunks = []

    def split_by_keyword(text, keyword):
        # Разделяем текст по ключевому слову и сохраняем ключевое слово в начале каждого чанка
        parts = re.split(f"({keyword})", text)
        chunks = [parts[i] + parts[i + 1] for i in range(1, len(parts) - 1, 2)]
        if parts[0]:
            chunks.insert(0, parts[0])
        if len(parts) % 2 == 0:
            chunks.append(parts[-1])
        return chunks

    if 'txt' in file_types and documents['txt'] is not None:
        for doc in documents['txt']:
            chunks = split_by_keyword(doc.page_content, keyword)
            for chunk in chunks:
                all_chunks.append(Document(source=doc.source, page_content=chunk, metadata=doc.metadata))

    if 'json' in file_types and documents['json'] is not None:
        for idx, doc in enumerate(documents['json']):
            text = json.dumps(doc, ensure_ascii=False)
            chunks = split_by_keyword(text, keyword)
            for chunk in chunks:
                all_chunks.append(Document(source=documents['json_metadata'][idx]['source'], page_content=chunk))

    if 'docx' in file_types and documents['docx'] is not None:
        for doc in documents['docx']:
            chunks = split_by_keyword(doc.page_content, keyword)
            for chunk in chunks:
                all_chunks.append(Document(source=doc.source, page_content=chunk, metadata=doc.metadata))

    return all_chunks


chunks_res = split_docs_to_chunks(DOCS, ['txt', 'json', 'docx'])


def get_chroma_vectorstore(documents, embeddings, persist_directory):
    if os.path.isdir(persist_directory) and os.listdir(persist_directory):
        print("Loading existing Chroma vectorstore...")
        vectorstore = Chroma(
            embedding_function=embeddings, persist_directory=persist_directory
        )

        existing_files = get_uploaded_filenames('local')
        uniq_sources_to_add = set(
            doc.metadata['source'] for doc in chunks_res
            if doc.metadata['source'] not in existing_files
        )

        if uniq_sources_to_add:
            vectorstore.add_documents(
                documents=[doc for doc in chunks_res if doc.metadata['source'] in uniq_sources_to_add],
                embedding=embeddings
            )
            for filename in uniq_sources_to_add:
                add_filename_to_metadata('local', filename)
        else:
            print('Новых документов не было, пропускаем шаг добавления')

    else:
        print("Creating and indexing new Chroma vectorstore...")
        vectorstore = Chroma.from_documents(
            documents=documents,
            embedding=embeddings, persist_directory=persist_directory
        )
        uniq_sources_to_add = set(doc.metadata['source'] for doc in documents)
        for filename in uniq_sources_to_add:
            add_filename_to_metadata('local', filename)

    return vectorstore


vectorstore = get_chroma_vectorstore(documents=chunks_res, embeddings=embeddings, persist_directory=CHROMA_PATH)

retriever = vectorstore.as_retriever(search_kwargs={"k": 2}, search_type='similarity')


def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)


chat_history_for_chain = SQLiteChatHistory()


prompt_sys = '''
Вы - ИИ-ассистент, специализирующийся на предоставлении информации о множестве автомобилей. Ваша задача - анализировать данные по всем автомобилям и отвечать на вопросы, основываясь на полном анализе всей доступной информации.

ВАЖНО: Вы СТРОГО ОГРАНИЧЕНЫ использованием ТОЛЬКО информации, предоставленной в следующем контексте:

Контекст:
{context}

СТРОГИЕ ПРАВИЛА:
1. Используйте ИСКЛЮЧИТЕЛЬНО информацию из предоставленного контекста. НИКОГДА не обращайтесь к внешним источникам или своим знаниям.
2. НЕ ВЫДУМЫВАЙТЕ информацию. Если ответа нет в контексте, скажите "Ответ не найден, пожалуйста, уточните ваш вопрос".
3. Контекст содержит несколько чанков, каждый чанк содержит информацию об одном автомобиле.
4. ВСЕГДА отвечайте на русском языке.

КРИТИЧЕСКИЕ ИНСТРУКЦИИ ПО ОБРАБОТКЕ ДАННЫХ:
1. ДЛЯ КАЖДОГО ВОПРОСА вы ОБЯЗАНЫ проанализировать ВСЕ чанки контекста, независимо от их кажущейся релевантности.
2. При ответе на вопросы о сравнении или поиске наилучшего/наихудшего значения характеристики (например, "у какого авто самый большой запас хода?"):
   a. Проверьте указанную характеристику в КАЖДОМ чанке.
   b. Создайте временный список всех найденных значений с указанием автомобиля.
   c. Сравните все значения между собой.
   d. Выберите наилучшее/наихудшее значение на основе сравнения.
   e. Укажите в ответе автомобиль с этим значением и само значение.
3. Если в каком-то чанке отсутствует запрашиваемая характеристика, отметьте это в ответе.
4. Всегда предоставляйте полный ответ, основанный на анализе ВСЕХ чанков.

ПРОЦЕСС ОТВЕТА:
1. Внимательно прочитайте вопрос пользователя.
2. Просмотрите и проанализируйте ВСЕ чанки контекста.
3. Если вопрос требует сравнения или поиска экстремального значения, выполните шаги, указанные в пункте 2 КРИТИЧЕСКИХ ИНСТРУКЦИЙ ПО ОБРАБОТКЕ ДАННЫХ.
4. Сформулируйте ваш ответ на основе полного анализа всех чанков.
5. Если информация отсутствует в каких-либо чанках, укажите это в ответе.

Для приветствия используйте: "Здравствуйте! Я готов помочь вам с информацией о нескольких автомобилях в нашем ассортименте. Чем могу быть полезен?"

Вопрос: {question}
'''


prompt_new = ChatPromptTemplate.from_messages(
    [
        (
            "system", prompt_sys,
        ),
        MessagesPlaceholder(variable_name="chat_history"),
        ("human", "{question}"),
    ]
)

chain_new = prompt_new | llm

chain_with_message_history = RunnableWithMessageHistory(
    chain_new,
    lambda session_id: chat_history_for_chain.messages(limit=15),
    input_messages_key="question",
    history_messages_key="chat_history",
)

app = FastAPI()

@app.get("/register", response_class=HTMLResponse)
async def get_register():
    return FileResponse("static/register.html")


def is_email_unique(email: str) -> bool:
    """Проверяет, является ли email уникальным в таблице Users."""
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT 1 FROM Users WHERE email = ?", (email,))
        return cursor.fetchone() is None




@app.get("/confirm-email")
async def confirm_email():
    try:
        with sqlite3.connect('metadata.db') as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id, email FROM Users ORDER BY id DESC LIMIT 1")
            user = cursor.fetchone()

            if user:
                user_id, username = user

                # Обновление статуса пользователя
                cursor.execute(
                    "UPDATE Users SET is_active = TRUE WHERE id = ?",
                    (user_id,)
                )
                conn.commit()

                # Логика создания сессии и добавления сообщения
                chat_history = SQLiteChatHistory(user_email=username, user_password="dummy_password")
                print(f"SQLiteChatHistory создан с user_email={username}")
                try:
                    session_id = chat_history.start_new_session()
                    if not check_session_id_exists(session_id):
                        print(f"Session started: {session_id}")
                        db_manager.add_chat_message(session_id,
                                                    "Вас приветствует На Связи! Напишите Ваш вопрос.",
                                                    "Система")
                    else:
                        print(f"Session started: {session_id}")
                    # Переадресация на клиентскую часть после активации
                    return RedirectResponse(url="/")
                except Exception as e:
                    print(f"Error starting session: {e}")
                    return JSONResponse(content={"status": "error", "message": "Failed to start session"},
                                        status_code=500)
            else:
                # return JSONResponse(content={"status": "error", "message": "Invalid or expired token."}, status_code=400)
                return FileResponse("static/error_page.html")

    except Exception as e:
        print(f"Error confirming email: {e}")
        raise HTTPException(status_code=500, detail="Internal Server Error")


@app.post("/register")
async def post_register(username: str = Form(...), password: str = Form(...)):
    if not is_email_unique(username):
        return JSONResponse(content={"status": "error", "message": "Пользователь с таким телефоном уже существует."},
                            status_code=401)
    else:
        user_id = add_user_to_db(username, password)
        await confirm_email()
        return JSONResponse(
            content={"status": "success", "message": "Вы в системе, задайте вопрос."},
            status_code=200)




# Функция для поиска по session_id в Chat
def check_session_id_exists(session_id: str) -> bool:
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT EXISTS(SELECT 1 FROM Chat WHERE session_id = ?)", (session_id,))
        return cursor.fetchone()[0] == 1




# Функция для поиска пользователя по email
def get_user_by_email(email: str):
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id, cyberman_id FROM Users WHERE email = ?", (email,))
        return cursor.fetchone()




# Функция для поиска сессии по user_id и cyberman_id
def get_session_by_user_and_cyberman(user_id: int, cyberman_id: int):
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM Session WHERE user_id = ? AND cyberman_id = ?", (user_id, cyberman_id))
        return cursor.fetchone()


@app.post("/create_new_chat/")
async def create_new_chat(request: Request):

    data = await request.json()
    email = data.get('email')

    if not email:
        raise HTTPException(status_code=400, detail="Email is required")

    user = get_user_by_email(email)
    if user is None:
        raise HTTPException(status_code=404, detail="User not found")

    user_id, cyberman_id = user
    session_id = db_manager.create_new_chat_session(user_id, cyberman_id)

    # Создаем начальное сообщение в чате с использованием нового session_id
    db_manager.add_chat_message(session_id, "Вас приветствует На Связи! Напишите Ваш вопрос.", "Система")

    return {"session_id": session_id}



@app.get("/get_chat_messages/{session_id}")
async def get_chat_messages(session_id: int):
    messages = db_manager.get_chat_messages_by_session_id(session_id)
    return {"messages": messages}


@app.get("/get_user_chats/{email}")
async def get_user_chats(email: str):
    user = get_user_by_email(email)
    if user is None:
        raise HTTPException(status_code=404, detail="User not found")

    user_id, _ = user
    chats = db_manager.get_chats_by_user_id(user_id)
    return {"chats": chats}



def get_session_history(session_id):
    history = chat_history_for_chain.messages(limit=15)
    print(f"Session {session_id} history: {history}")
    return history



@app.websocket("/ws/rag_chat/")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()

    email = websocket.query_params.get('email')
    if email is None:
        await websocket.send_json({"error": "Требуется email"})
        await websocket.close(code=status.WS_1008_POLICY_VIOLATION)
        return

    user = get_user_by_email(email)
    if user is None:
        await websocket.send_json({"error": "Пользователь не найден"})
        await websocket.close(code=status.WS_1008_POLICY_VIOLATION)
        return

    user_id, cyberman_id = user

    session = get_session_by_user_and_cyberman(user_id, cyberman_id)
    if session is None:
        await websocket.send_json({"error": "Сессия не найдена"})
        await websocket.close(code=status.WS_1008_POLICY_VIOLATION)
        return

    session_id = session[0]
    chat_history_for_chain.current_session_id = session_id

    messages = db_manager.get_chat_messages_by_session_id(session_id)
    await websocket.send_json({"messages": messages})

    try:
        while True:
            data = await websocket.receive_json()
            question_data = data.get('question_data')
            if question_data is None:
                await websocket.send_json({"error": "Требуется question_data"})
                continue

            question = question_data.get('question')
            new_session_id = question_data.get('session_id')

            if new_session_id:
                session_id = new_session_id
                chat_history_for_chain.current_session_id = session_id  # Обновляем session_id

            if question is None:
                await websocket.send_json({"error": "Требуется question"})
                continue

            context = format_docs(retriever.invoke(question))
            history = get_session_history(session_id)

            # print(f"History: {history}")

            request_payload = {
                "question": question,
                "context": context,
                "chat_history": history
            }
            print(f"Request payload: {request_payload}")

            try:
                # answer = chain_with_message_history.invoke(
                #     {"question": question, "context": format_docs(retriever.invoke(question))},
                #     {"configurable": {"session_id": session_id}}
                # ).content
                answer = chain_with_message_history.invoke(
                    request_payload,
                    {"configurable": {"session_id": session_id}}
                ).content

            except Exception as e:
                await websocket.send_json({"error": str(e)})
                continue

            if answer:
                chat_history_for_chain.add_message(HumanMessage(content=question))
                chat_history_for_chain.add_message(AIMessage(content=answer))

            db_manager.add_chat_message(session_id, question, "human")
            db_manager.add_chat_message(session_id, answer, "ai")

            messages = db_manager.get_chat_messages_by_session_id(session_id)


            await websocket.send_json({"answer": answer})

    except WebSocketDisconnect:
        chat_history_for_chain.end_session()



from fastapi.middleware.cors import CORSMiddleware



# Добавляем middleware для обработки CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],  # Разрешаем конкретный фронтенд
    allow_credentials=True,
    allow_methods=["*"],  # Разрешаем любые HTTP методы
    allow_headers=["*"],  # Разрешаем любые заголовки
)




DATABASE_URL = "sqlite:///./metadata.db"
engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


@app.post("/save_chat_history/{session_id}")
async def save_chat_history(session_id: int, request: Request, db: Session = Depends(get_db)):
    try:
        # Получаем данные из запроса
        data = await request.json()
        username = data.get('username', 'Unknown User')  # Получаем username или используем значение по умолчанию

        # Получение сообщений из таблицы Chat
        chat_query = text("""
            SELECT messages, sender
            FROM Chat
            WHERE session_id = :session_id
            ORDER BY sent_at
        """)
        chat_result = db.execute(chat_query, {"session_id": session_id}).fetchall()

        human_messages = [message for message, sender in chat_result if sender == 'human']

        if len(human_messages) < 2:
            return {"message": "Not enough human messages to send data to Bitrix24"}

        # Форматирование диалога
        dialog = []
        for message, sender in chat_result:
            dialog.append(f"{sender}: {message}")

        # Подготовка данных для отправки в Bitrix24
        bitrix_data = {
            "fields": {
                "TITLE": f"Лид от {username} (Session {session_id})",  # Используем username в TITLE
                "NAME": f"{username} (Session {session_id})",
                "COMMENTS": "\n\n".join(dialog)
            },
            "params": {
                "REGISTER_SONET_EVENT": "Y"
            }
        }

        # Отправка данных в Bitrix24
        bitrix_url = "https://b24-d61yk4.bitrix24.by/rest/8/out59101i3tfsf77/crm.lead.add.json"  # Замените на ваш вебхук URL
        response = requests.post(bitrix_url, json=bitrix_data)

        if response.status_code != 200:
            raise HTTPException(status_code=500, detail="Failed to send data to Bitrix24")

        return {"message": "Chat history saved and sent to Bitrix24 successfully"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))




if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8222)
